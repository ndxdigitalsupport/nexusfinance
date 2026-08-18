"""
NexusFinance document engine — professional Word template generator.

Design language:
  - Brand teal accent (#00BDAA), deep navy ink (#0B1F2E), neutral greys
  - Clean sans body (Segoe UI / Khmer OS Siemreap), geometric cover band
  - Styled headings with accent rule, header/footer chrome, page numbers
  - Consistent tables, callout boxes, code blocks, and lists

Usage (EN and KH):
    from engine import NexusDoc
    doc = NexusDoc(lang="EN", title="...", subtitle="...", doc_number="...")
    doc.cover(version="1.0", date="2026-08-18")
    doc.h1("Section")
    doc.p("Body text")
    doc.table(["A","B"], [["1","2"]])
    doc.save("output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import re

# ── Brand palette ────────────────────────────────────────────
TEAL        = RGBColor(0x00, 0xBD, 0xAA)
TEAL_DARK   = RGBColor(0x00, 0x8A, 0x7C)
TEAL_SOFT   = RGBColor(0xE0, 0xF7, 0xF4)
NAVY        = RGBColor(0x0B, 0x1F, 0x2E)
NAVY_SOFT   = RGBColor(0x14, 0x33, 0x48)
SLATE       = RGBColor(0x47, 0x5A, 0x6B)
GREY        = RGBColor(0x6B, 0x7A, 0x89)
LIGHT       = RGBColor(0xED, 0xF1, 0xF5)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GOLD        = RGBColor(0xC9, 0xA2, 0x5C)

# ── Fonts ────────────────────────────────────────────────────
FONT_EN = "Segoe UI"
FONT_KH = "Khmer OS Siemreap"
FONT_MONO = "Consolas"

class NexusDoc:
    def __init__(self, lang="EN", title="", subtitle="", doc_number="", footer_note="Internal & Confidential"):
        self.lang = lang
        self.body_font = FONT_EN if lang == "EN" else FONT_KH
        self.title = title
        self.subtitle = subtitle
        self.doc_number = doc_number or "NF-DOC"
        self.footer_note = footer_note
        self.doc = Document()

        # Page geometry
        sec = self.doc.sections[0]
        sec.page_width = Inches(8.27)   # A4
        sec.page_height = Inches(11.69)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)

        self._setup_styles()

    # ── Base styles ──────────────────────────────────────────
    def _setup_styles(self):
        st = self.doc.styles["Normal"]
        st.font.name = self.body_font
        st.font.size = Pt(10.5)
        st.font.color.rgb = NAVY
        pf = st.paragraph_format
        pf.space_after = Pt(7)
        pf.line_spacing = 1.35

        # East Asian / complex script font binding
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:cs"), FONT_KH if self.lang == "KH" else FONT_EN)

    def _set_run_font(self, run, size=10.5, bold=False, color=NAVY, font=None, italic=False):
        run.font.name = font or self.body_font
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:cs"), FONT_KH if self.lang == "KH" else FONT_EN)
        return run

    # ── Low-level helpers ────────────────────────────────────
    @staticmethod
    def _shade_cell(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    @staticmethod
    def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for tag, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
            el = OxmlElement(f"w:{tag}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def _page_break(self):
        self.doc.add_page_break()

    # ── Paragraph builders ───────────────────────────────────
    def p(self, text, size=10.5, color=NAVY, bold=False, italic=False,
          align=None, space_before=None, space_after=None, font=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        self._set_run_font(run, size=size, bold=bold, color=color, italic=italic, font=font)
        if align: para.alignment = align
        if space_before is not None: para.paragraph_format.space_before = Pt(space_before)
        if space_after is not None: para.paragraph_format.space_after = Pt(space_after)
        return para

    def rich(self, segments, size=10.5, space_after=7):
        """segments: list of (text, dict) with keys bold, italic, color, font"""
        para = self.doc.add_paragraph()
        for text, fmt in segments:
            run = para.add_run(text)
            self._set_run_font(
                run,
                size=fmt.get("size", size),
                bold=fmt.get("bold", False),
                italic=fmt.get("italic", False),
                color=fmt.get("color", NAVY),
                font=fmt.get("font"),
            )
        para.paragraph_format.space_after = Pt(space_after)
        return para

    def bullet(self, text, level=0, size=10.5, color=NAVY):
        para = self.doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        run = para.add_run(text)
        self._set_run_font(run, size=size, color=color)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Inches(0.3 + level * 0.28)
        return para

    def numbered(self, text, size=10.5):
        para = self.doc.add_paragraph(style="List Number")
        run = para.add_run(text)
        self._set_run_font(run, size=size)
        para.paragraph_format.space_after = Pt(3)
        return para

    def check_item(self, text, checked=True, size=10.5):
        mark = "\u2611  " if checked else "\u2610  "
        para = self.doc.add_paragraph()
        run_mark = para.add_run(mark)
        self._set_run_font(run_mark, size=size, bold=True, color=TEAL_DARK)
        run = para.add_run(text)
        self._set_run_font(run, size=size)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.15)
        return para

    # ── Headings (use Word built-in styles for TOC support) ──
    def _heading(self, text, level, size, color):
        # Use Word's built-in heading styles so the TOC field can find them
        para = self.doc.add_heading(text, level=level)
        # Override all runs with brand formatting
        for run in para.runs:
            self._set_run_font(run, size=size, bold=True, color=color, font=self.body_font)
        if level == 1:
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
            # accent rule under H1
            pPr = para._element.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), "00BDAA")
            pbdr.append(bottom)
            pPr.append(pbdr)
        else:
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
        return para

    def h1(self, text):
        return self._heading(text, 1, 16, NAVY)

    def h2(self, text):
        return self._heading(text, 2, 13, TEAL_DARK)

    def h3(self, text):
        return self._heading(text, 3, 11.5, SLATE)

    # ── Tables ───────────────────────────────────────────────
    def table(self, headers, rows, col_widths=None, font_size=9.5, header_fill="0B1F2E"):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        t.autofit = False

        # header
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(h)
            self._set_run_font(run, size=font_size, bold=True, color=WHITE, font=self.body_font)
            self._shade_cell(cell, header_fill)
            self._set_cell_margins(cell)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # body rows
        for ri, row in enumerate(rows):
            fill = "F4F8FA" if ri % 2 == 1 else None
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(val))
                self._set_run_font(run, size=font_size, color=NAVY, font=self.body_font)
                if fill:
                    self._shade_cell(cell, fill)
                self._set_cell_margins(cell)
                para.paragraph_format.space_after = Pt(1)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # column widths
        if col_widths:
            for row in t.rows:
                for ci, w in enumerate(col_widths):
                    row.cells[ci].width = Inches(w)

        # spacer
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ── Callouts ─────────────────────────────────────────────
    def callout(self, title, text, kind="info"):
        colors = {
            "info":   ("00BDAA", "E0F7F4"),
            "warn":   ("C9A25C", "FBF3E0"),
            "danger": ("D9534F", "FCE9E8"),
            "tip":    ("2E7D32", "E8F5E9"),
        }
        border_hex, fill_hex = colors.get(kind, colors["info"])

        t = self.doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        cell = t.rows[0].cells[0]
        cell.text = ""

        # accent left border
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "28")
        left.set(qn("w:space"), "0")
        left.set(qn("w:color"), border_hex)
        tcBorders.append(left)
        tcPr.append(tcBorders)
        self._shade_cell(cell, fill_hex)
        self._set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        if title:
            p = cell.paragraphs[0]
            run = p.add_run(title)
            self._set_run_font(run, size=10.5, bold=True, color=NAVY)
            p.paragraph_format.space_after = Pt(3)
        p2 = cell.add_paragraph()
        run2 = p2.add_run(text)
        self._set_run_font(run2, size=10, color=NAVY)
        p2.paragraph_format.space_after = Pt(0)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ── Code block ───────────────────────────────────────────
    def code(self, text):
        t = self.doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        cell = t.rows[0].cells[0]
        cell.text = ""
        self._shade_cell(cell, "0B1F2E")
        self._set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        for i, line in enumerate(text.split("\n")):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            run = para.add_run(line)
            self._set_run_font(run, size=9, color=RGBColor(0xBF, 0xE8, 0xE2), font=FONT_MONO)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.15
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ── Divider ──────────────────────────────────────────────
    def divider(self):
        para = self.doc.add_paragraph()
        pPr = para._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "D8DEE4")
        pbdr.append(bottom)
        pPr.append(pbdr)
        para.paragraph_format.space_after = Pt(6)
        return para

    # ── Cover page ───────────────────────────────────────────
    def cover(self, version="1.0", date="2026", author="NexusFinance", tags=None):
        doc = self.doc
        sec = doc.sections[0]
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
        hdr = sec.header
        for p in list(hdr.paragraphs):
            p.text = ""
        hdr_p = hdr.paragraphs[0]
        hdr_p.text = ""
        ftr = sec.footer
        for p in list(ftr.paragraphs):
            p.text = ""
        ftr_p = ftr.paragraphs[0]
        ftr_p.text = ""
        run = ftr_p.add_run(f"{self.title}  •  {self.doc_number}  •  {self.footer_note}")
        self._set_run_font(run, size=8, color=GREY)
        ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # top brand band
        band = doc.add_table(rows=1, cols=1)
        band.style = "Table Grid"
        bandcell = band.rows[0].cells[0]
        bandcell.text = ""
        self._shade_cell(bandcell, "0B1F2E")
        self._set_cell_margins(bandcell, top=180, bottom=180, left=240, right=240)
        bp = bandcell.paragraphs[0]
        brun = bp.add_run("NEXUSFINANCE")
        self._set_run_font(brun, size=13, bold=True, color=TEAL)
        bp2 = bandcell.add_paragraph()
        brun2 = bp2.add_run("Secure • Modern • Cutting-edge  |  Lending Platform for Cambodia")
        self._set_run_font(brun2, size=9, color=RGBColor(0x9F, 0xB4, 0xC4))
        bp2.paragraph_format.space_after = Pt(0)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # spacer
        doc.add_paragraph().paragraph_format.space_after = Pt(30)

        # title
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        trun = tp.add_run(self.title)
        self._set_run_font(trun, size=34, bold=True, color=NAVY)
        tp.paragraph_format.space_after = Pt(6)

        # accent rule under title
        rule = doc.add_paragraph()
        pPr = rule._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "24")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "00BDAA")
        pbdr.append(bottom)
        pPr.append(pbdr)
        rule.paragraph_format.space_after = Pt(12)

        # subtitle
        sp = doc.add_paragraph()
        srun = sp.add_run(self.subtitle)
        self._set_run_font(srun, size=13, color=SLATE)
        sp.paragraph_format.space_after = Pt(4)

        # accent line below subtitle
        accent = doc.add_paragraph()
        arun = accent.add_run(self.lang)
        self._set_run_font(arun, size=10, bold=True, color=TEAL_DARK)
        accent.paragraph_format.space_after = Pt(0)

        doc.add_paragraph().paragraph_format.space_after = Pt(30)

        # meta table
        mt = doc.add_table(rows=0, cols=2)
        mt.style = "Table Grid"
        meta_rows = [
            ("Document Number", self.doc_number),
            ("Version", version),
            ("Date", date),
            ("Prepared By", author),
            ("Classification", self.footer_note),
        ]
        for k, v in meta_rows:
            row = mt.add_row()
            c0, c1 = row.cells
            c0.text = ""
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(k)
            self._set_run_font(r0, size=10, bold=True, color=NAVY)
            self._shade_cell(c0, "F4F8FA")
            self._set_cell_margins(c0)
            c1.text = ""
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(v)
            self._set_run_font(r1, size=10, color=SLATE)
            self._set_cell_margins(c1)
        # width fix
        for row in mt.rows:
            row.cells[0].width = Inches(2.2)
            row.cells[1].width = Inches(4.2)

        if tags:
            doc.add_paragraph().paragraph_format.space_after = Pt(14)
            tagline = doc.add_paragraph()
            tr = tagline.add_run("  •  ".join(tags))
            self._set_run_font(tr, size=9, color=GREY)

        doc.add_page_break()
        return doc

    # ── TOC placeholder ──────────────────────────────────────
    def toc(self, entries):
        """Build a manual Table of Contents.
        entries: list of (section_number, title) e.g. [("1", "Executive Summary"), ("2", "Architecture")]
        """
        title_text = "Table of Contents" if self.lang == "EN" else "មាតិកា"
        p = self.doc.add_paragraph()
        r = p.add_run(title_text)
        self._set_run_font(r, size=16, bold=True, color=NAVY)
        p.paragraph_format.space_after = Pt(6)

        # accent rule under title
        pPr = p._element.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "00BDAA")
        pbdr.append(bottom)
        pPr.append(pbdr)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(8)

        for num, title in entries:
            toc_para = self.doc.add_paragraph()
            toc_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=2, leader=2)
            toc_para.paragraph_format.space_after = Pt(4)

            num_run = toc_para.add_run(f"{num}.  ")
            self._set_run_font(num_run, size=11, bold=True, color=TEAL_DARK, font=self.body_font)

            title_run = toc_para.add_run(title)
            self._set_run_font(title_run, size=11, bold=False, color=NAVY, font=self.body_font)

            # add tab + dots (Word handles leader via tab stop)

        self.doc.add_page_break()
        return self.doc

    # ── Document control / revision table ────────────────────
    def revision_table(self, rows):
        self.h2("Document Control" if self.lang == "EN" else "ការត្រួតពិនិត្យឯកសារ")
        self.table(
            ["Version", "Date", "Author", "Description"] if self.lang == "EN"
            else ["កំណែ", "កាលបរិច្ឆេទ", "អ្នកសរសេរ", "ការពិពណ៌នា"],
            rows,
            col_widths=[0.9, 1.1, 1.6, 2.8],
        )

    # ── Footer page numbers on all sections ──────────────────
    def add_page_numbers(self):
        for section in self.doc.sections:
            ftr = section.footer
            p = ftr.paragraphs[0]
            # add "Page X of Y"
            r = p.add_run("  •  Page ")
            self._set_run_font(r, size=8, color=GREY)
            # PAGE field
            fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
            fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
            t = OxmlElement("w:t"); t.text = "1"
            fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
            run_el = p.add_run()._r
            run_el.append(fld1); run_el.append(instr); run_el.append(fld2); run_el.append(t); run_el.append(fld3)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return self.doc

    # ── Save ─────────────────────────────────────────────────
    def save(self, path):
        self.add_page_numbers()
        self.doc.save(path)
        try:
            print(f"  [OK] saved: {path}")
        except UnicodeEncodeError:
            safe = path.encode("ascii", "replace").decode("ascii")
            print(f"  [OK] saved: {safe}")
