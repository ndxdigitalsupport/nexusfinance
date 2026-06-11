from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helpers ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return h

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ══════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KHQR PAYMENT INTEGRATION GUIDE')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Bakong (NBC) + ABA Bank Integration for NexusFinance')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph('─' * 60)

# ══════════════════════════════════════════════════
# PART 1: BAKONG
# ══════════════════════════════════════════════════
add_heading_styled('Part 1: Bakong (NBC) KHQR Integration', level=1)

add_heading_styled('Overview', level=2)
doc.add_paragraph('Bakong is Cambodia\'s national real-time payment system operated by the National Bank of Cambodia (NBC). KHQR is the standardized QR code format used by all Cambodian banks.')

doc.add_paragraph('How it works:')
add_bullet('Customer scans KHQR code')
add_bullet('Opens Bakong or any banking app')
add_bullet('Confirms payment')
add_bullet('NBC clears funds')
add_bullet('Webhook notifies your server')

add_heading_styled('Prerequisites', level=2)
add_table(
    ['Item', 'How to Get'],
    [
        ['Bakong Wallet', 'Download Bakong app, register'],
        ['Bakong Account ID', 'Format: yourname@bank (e.g. merchant@aclb)'],
        ['NBC API Token', 'Register with NBC for developer access'],
        ['Server in Cambodia', 'Production check-transaction API only works from Cambodia IPs'],
    ]
)

add_heading_styled('Getting an NBC API Token', level=2)
add_bullet('Email NBC at their developer support or contact your bank')
add_bullet('They provide a username/password for the API')
add_bullet('Call POST /v1/request_token with your credentials')
add_bullet('Receive a Bearer token (expires - use /v1/renew_token)')

add_heading_styled('API Reference', level=2)

add_heading_styled('Base URLs', level=3)
add_table(
    ['Environment', 'URL'],
    [
        ['Sandbox (SIT)', 'https://sit-api-bakong.nbc.gov.kh/v1'],
        ['Production', 'https://api-bakong.nbc.gov.kh/v1'],
    ]
)

add_heading_styled('Endpoints', level=3)
add_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/v1/request_token', 'Get access token'],
        ['POST', '/v1/verify', 'Verify token validity'],
        ['POST', '/v1/renew_token', 'Renew expiring token'],
        ['POST', '/v1/generate_deeplink_by_qr', 'Create payment deeplink'],
        ['POST', '/v1/check_transaction_by_md5', 'Check payment by MD5'],
        ['POST', '/v1/check_transaction_by_hash', 'Check payment by full hash'],
        ['POST', '/v1/check_transaction_by_short_hash', 'Check payment by short hash'],
        ['POST', '/v1/check_bakong_account', 'Verify Bakong account exists'],
    ]
)

add_heading_styled('1. Request Token', level=3)
add_code('POST {{baseUrl}}/v1/request_token\n\nBody:\n{\n  "username": "your_registered_email@example.com",\n  "password": "your_api_password"\n}\n\nResponse:\n{\n  "responseCode": 0,\n  "message": "Success",\n  "data": {\n    "token": "eyJhbGciOiJIUzI1NiIs..."\n  }\n}')

add_heading_styled('2. Generate Deeplink', level=3)
add_code('POST {{baseUrl}}/v1/generate_deeplink_by_qr\nAuthorization: Bearer {{token}}\n\nBody:\n{\n  "qr": "000201010212...full_khqr_string",\n  "md5": "e10adc3949ba59abbe56e057f20f883e"\n}')

add_heading_styled('3. Check Transaction by MD5', level=3)
add_code('POST {{baseUrl}}/v1/check_transaction_by_md5\nAuthorization: Bearer {{token}}\n\nBody:\n{\n  "md5": "e10adc3949ba59abbe56e057f20f883e"\n}\n\nResponse:\n{\n  "responseCode": 0,\n  "message": "Success",\n  "data": {\n    "status": "SUCCESS",\n    "amount": 25.00,\n    "currency": "USD",\n    "hash": "9bc29426b5e55f171b0e137478d04eb8",\n    "fromAccountId": "customer@aclb",\n    "toAccountId": "merchant@aclb",\n    "transactionDate": "2025-01-15T10:30:00+07:00"\n  }\n}')

add_heading_styled('Response Codes', level=3)
add_table(
    ['Code', 'Meaning'],
    [
        ['0', 'Success'],
        ['1', 'Fail / Not found'],
        ['2', 'Static QR not supported'],
        ['3', 'Transaction failed'],
        ['4', 'Deeplink provider error'],
        ['5', 'Missing required fields'],
        ['6', 'Unauthorized'],
    ]
)

add_heading_styled('KHQR String Format (EMVCo)', level=2)
doc.add_paragraph('KHQR follows the EMVCo QR standard with these tags:')
add_table(
    ['Tag', 'Field', 'Example'],
    [
        ['00', 'Payload Format Indicator', '01'],
        ['01', 'Point of Initiation Method', '11 (static) or 12 (dynamic)'],
        ['29', 'Bakong Account ID', 'merchant@aclb'],
        ['53', 'Currency Code', '840 (USD) or 116 (KHR)'],
        ['54', 'Amount', '25.00'],
        ['58', 'Country Code', 'KH'],
        ['59', 'Merchant Name', 'NexusFinance'],
        ['60', 'Merchant City', 'Phnom Penh'],
        ['61', 'Store Label', 'Main Branch'],
        ['63', 'CRC', '7D00 (CRC16-CCITT)'],
    ]
)

add_heading_styled('Reference Links', level=2)
add_bullet('Bakong API Docs: https://bakong.nbc.gov.kh/download/KHQR/integration/Bakong%20Open%20API%20Document.pdf')
add_bullet('Bakong Production API: https://api-bakong.nbc.gov.kh/v1')
add_bullet('Bakong Sandbox API: https://sit-api-bakong.nbc.gov.kh/v1')
add_bullet('KHQR SDK (Java): https://github.com/tongbora/Bakong-API-Integration-with-Spring-Boot')
add_bullet('KHQR Integration (Node.js): https://github.com/Leap-Chanvuthy/bakong-open-api')
add_bullet('CamFinTech Guide: https://www.camfintech.com/knowledge/bakong-technical-integration')

# ══════════════════════════════════════════════════
# PART 2: ABA
# ══════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('Part 2: ABA Bank API Integration', level=1)

add_heading_styled('Overview', level=2)
doc.add_paragraph('ABA Bank offers a separate merchant API for payment collection. Unlike Bakong (which is interbank), the ABA API is specific to ABA accounts only.')

add_heading_styled('Key Differences from Bakong', level=2)
add_table(
    ['Feature', 'Bakong', 'ABA'],
    [
        ['Coverage', 'All Cambodian banks', 'ABA accounts only'],
        ['QR Standard', 'KHQR (EMVCo)', 'ABA-specific QR'],
        ['API Access', 'Apply via NBC', 'Apply via ABA'],
        ['Settlement', 'Real-time via NBC', 'Batch settlement'],
    ]
)

add_heading_styled('Prerequisites', level=2)
add_table(
    ['Item', 'How to Get'],
    [
        ['ABA Merchant Account', 'Apply at any ABA branch'],
        ['ABA Corporate Account', 'Business account with ABA'],
        ['API Credentials', 'Contact ABA Business Banking'],
        ['Merchant ID', 'Provided by ABA upon approval'],
        ['API Key/Secret', 'Provided by ABA upon approval'],
        ['IP Whitelist', 'Register your server IP with ABA'],
    ]
)

add_heading_styled('How to Apply', level=2)
doc.add_paragraph('1. Visit any ABA Bank branch with:')
add_bullet('Business registration / company license')
add_bullet('ID documents')
add_bullet('Bank account statement')
doc.add_paragraph('2. Request "Merchant Payment Gateway / API integration"')
doc.add_paragraph('3. ABA provides:')
add_bullet('Merchant ID')
add_bullet('API Key & Secret')
add_bullet('Endpoint URLs')
add_bullet('IP whitelist registration')

add_heading_styled('Common Endpoints', level=2)
add_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/api/v1/payment/init', 'Create payment request'],
        ['POST', '/api/v1/payment/status', 'Check payment status'],
        ['POST', '/api/v1/payment/refund', 'Process refund'],
        ['POST', '/api/v1/account/balance', 'Check account balance'],
    ]
)

add_heading_styled('Initiate Payment', level=2)
add_code('POST {{abaBaseUrl}}/api/v1/payment/init\n\nHeaders:\nContent-Type: application/json\nX-Merchant-Id: {{merchantId}}\nX-Signature: {{hmac_sha256_signature}}\nX-Timestamp: 2025-01-15T10:00:00Z\n\nBody:\n{\n  "txnId": "ORDER_20250115_001",\n  "amount": 25.00,\n  "currency": "USD",\n  "description": "Loan repayment",\n  "callbackUrl": "https://nexusfinancefintech.vercel.app/api/payment/callback",\n  "customerPhone": "85512345678",\n  "customerEmail": "customer@example.com"\n}\n\nResponse:\n{\n  "responseCode": 0,\n  "message": "Success",\n  "data": {\n    "paymentUrl": "https://pay.aba.com.kh/checkout/ORDER_20250115_001",\n    "qrCode": "data:image/png;base64,...",\n    "expiresAt": "2025-01-15T10:30:00Z"\n  }\n}')

add_heading_styled('Payment Callback (Webhook)', level=2)
doc.add_paragraph('ABA sends a POST to your callbackUrl when payment completes:')
add_code('{\n  "txnId": "ORDER_20250115_001",\n  "status": "SUCCESS",\n  "amount": 25.00,\n  "currency": "USD",\n  "abaTxnRef": "ABA123456789",\n  "settlementDate": "2025-01-16",\n  "signature": "hmac_sha256_signature_for_verification"\n}')

# ══════════════════════════════════════════════════
# PART 3: CURRENT STATUS
# ══════════════════════════════════════════════════
doc.add_page_break()
add_heading_styled('Part 3: Current Implementation Status', level=1)

add_heading_styled('What Currently Exists in This Project', level=2)
add_table(
    ['File', 'Status', 'Description'],
    [
        ['server/khqr.ts', 'Mock', 'EMVCo QR generation, CRC verify, decode, deeplink, transaction check'],
        ['src/components/KHQRPage.tsx', 'UI', 'Full frontend with 5 tabs (Generate, Verify, Decode, Deeplink, Check)'],
        ['server/index.ts', 'Routes', '/api/khqr/* endpoints wired'],
    ]
)

add_heading_styled('What is Mock vs Real', level=2)
add_table(
    ['Feature', 'Current (Mock)', 'Real Implementation'],
    [
        ['QR Generation', 'Local EMVCo builder', 'NBC KHQR SDK'],
        ['Payment Verification', 'Random (30% success)', 'Call Bakong API'],
        ['Deeplink', 'Generates fake bakong:// link', 'Call Bakong generate_deeplink'],
        ['Transaction Storage', 'In-memory Map', 'Database'],
        ['Bank Coverage', 'Bakong only', 'Bakong + ABA'],
    ]
)

add_heading_styled('What Needs to Change for Production', level=2)
add_bullet('Replace server/khqr.ts mocks with real API calls to Bakong')
add_bullet('Add ABA API module (server/aba.ts)')
add_bullet('Store transactions in database instead of in-memory')
add_bullet('Add webhook endpoints for payment callbacks')
add_bullet('Add environment variables for API credentials')

# ══════════════════════════════════════════════════
# PART 4: ENV VARS
# ══════════════════════════════════════════════════
add_heading_styled('Part 4: Environment Variables', level=1)

add_heading_styled('Bakong (NBC)', level=2)
add_code('BAKONG_API_BASE_URL=https://sit-api-bakong.nbc.gov.kh/v1\nBAKONG_USERNAME=your_email@example.com\nBAKONG_PASSWORD=your_api_password\nBAKONG_MERCHANT_ID=merchant@aclb\nBAKONG_MERCHANT_NAME=NexusFinance\nBAKONG_MERCHANT_CITY=Phnom Penh')

add_heading_styled('ABA Bank', level=2)
add_code('ABA_API_BASE_URL=https://api.aba.com.kh/v1\nABA_MERCHANT_ID=your_aba_merchant_id\nABA_API_KEY=your_aba_api_key\nABA_API_SECRET=your_aba_api_secret\nABA_CALLBACK_URL=https://nexusfinancefintech.vercel.app/api/payment/callback')

# ══════════════════════════════════════════════════
# PART 5: CHECKLIST
# ══════════════════════════════════════════════════
add_heading_styled('Part 5: Quick Start Checklist', level=1)

checklist_items = [
    'Download Bakong app and register a wallet',
    'Get your Bakong Account ID (e.g. yourname@aclb)',
    'Email NBC or your bank to request API token',
    'Set BAKONG_API_* variables in .env',
    'Update server/khqr.ts to call real Bakong API instead of mock',
    'Visit ABA branch with business docs to apply for merchant API',
    'Set ABA_API_* variables in .env',
    'Register server IP with ABA (for Render: contact Render for static IP)',
    'Create server/aba.ts for ABA payment flow',
    'Add webhook endpoint /api/payment/callback',
    'Test in sandbox environment',
    'Go live in production',
]

for item in checklist_items:
    add_bullet('[ ] ' + item)

# ── Save ──
output_path = 'C:\\Users\\Asus\\Desktop\\yoooo\\KHQR_Integration_Guide.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
